import os
os.environ['KMP_DUPLICATE_LIB_OK']='TRUE'

import fitz
from docx import Document
from docx.shared import Inches
import win32com.client
import pythoncom
import pandas as pd
import tempfile
import io
import subprocess
import easyocr
import moviepy.video.io.VideoFileClip as VideoClip

def convert_pdf_to_docx(pdf_bytes):
    """Modified to work with in-memory files"""
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as pdf_temp:
        pdf_temp.write(pdf_bytes)
        pdf_temp_path = pdf_temp.name

    doc = Document()
    pdf_document = fitz.open(pdf_temp_path)
    
    try:
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            doc.add_paragraph(f"Page {page_num + 1}", style='Heading 1')
            text = page.get_text("text")
            if text.strip():
                doc.add_paragraph(text)

            images = page.get_images(full=True)
            for img_index, img in enumerate(images):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{base_image['ext']}") as img_temp:
                    img_temp.write(image_bytes)
                    img_path = img_temp.name
                
                try:
                    doc.add_picture(img_path, width=Inches(5.0))
                finally:
                    try:
                        os.unlink(img_path)
                    except:
                        pass

        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes.getvalue()

    finally:
        pdf_document.close()
        try:
            os.unlink(pdf_temp_path)
        except:
            pass

def convert_docx_to_pdf(docx_bytes):
    """Modified to work with in-memory files"""
    pythoncom.CoInitialize()
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as docx_temp:
        docx_temp.write(docx_bytes)
        docx_path = docx_temp.name

    pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
    
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        try:
            doc = word.Documents.Open(docx_path)
            doc.SaveAs(pdf_path, FileFormat=17)
            doc.Close()
            
            with open(pdf_path, 'rb') as pdf_file:
                pdf_bytes = pdf_file.read()
            
            return pdf_bytes
        
        finally:
            word.Quit()
            subprocess.run(['taskkill', '/F', '/IM', 'WINWORD.EXE'], 
                         stderr=subprocess.DEVNULL, stdout=subprocess.DEVNULL)
    
    finally:
        pythoncom.CoUninitialize()
        try:
            os.unlink(docx_path)
        except:
            pass
        try:
            os.unlink(pdf_path)
        except:
            pass

def convert_excel_to_csv(excel_bytes):
    """Modified to work with in-memory files"""
    df = pd.read_excel(excel_bytes)
    csv_buffer = io.StringIO()
    df.to_csv(csv_buffer, index=False)
    return csv_buffer.getvalue().encode()

def convert_csv_to_excel(csv_bytes):
    """Modified to work with in-memory files"""
    df = pd.read_csv(io.BytesIO(csv_bytes))
    excel_buffer = io.BytesIO()
    df.to_excel(excel_buffer, index=False)
    return excel_buffer.getvalue()

def convert_image_to_text(image_bytes):
    """Convert image bytes to text using EasyOCR"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as img_temp:
            img_temp.write(image_bytes)
            img_temp_path = img_temp.name

        reader = easyocr.Reader(['en'])
        
        results = reader.readtext(img_temp_path)
        
        text = '\n'.join([result[1] for result in results])
        
        return text.encode('utf-8')
    
    finally:
        try:
            os.unlink(img_temp_path)
        except:
            pass

def get_video_codec(format):
    """Get the appropriate codec for the target format"""
    codec_map = {
        'mp4': 'libx264',
        'avi': 'libxvid',
        'mkv': 'libx264',
        'mov': 'libx264',
        'wmv': 'libx264',
        'flv': 'flv'
    }
    return codec_map.get(format, 'libx264')

def convert_video_bytes(video_bytes, target_format):
    """Convert video bytes to another format"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4') as vid_temp:
            vid_temp.write(video_bytes)
            input_path = vid_temp.name
            
        output_path = os.path.splitext(input_path)[0] + f'.{target_format}'
        
        codec = get_video_codec(target_format)
        
        video = VideoClip.VideoFileClip(input_path)
        video.write_videofile(output_path, codec=codec)
        video.close()
        
        with open(output_path, 'rb') as f:
            converted_bytes = f.read()
            
        return converted_bytes
    
    finally:
        try:
            os.unlink(input_path)
            os.unlink(output_path)
        except:
            pass

def convert_file(file_bytes, input_format, target_format):
    """
    Convert file from one format to another
    
    Args:
        file_bytes (bytes): Input file as bytes
        input_format (str): Input file format (extension without dot)
        target_format (str): Target file format (extension without dot)
    
    Returns:
        tuple: (bytes, str) - Converted file bytes and MIME type
    """
    
    try:
        if len(file_bytes) == 0:
            raise ValueError("Empty file provided")

        # Video conversions
        if input_format in ['mp4', 'avi', 'mov', 'mkv', 'wmv', 'flv']:
            if target_format in ['mp4', 'avi', 'mkv', 'mov', 'wmv', 'flv']:
                converted_bytes = convert_video_bytes(file_bytes, target_format)
                mime_type = f"video/{target_format}"
            else:
                raise ValueError(f"Unsupported video target format: {target_format}")

        # Image to text conversion
        elif input_format in ['png', 'jpg', 'jpeg'] and target_format == 'txt':
            converted_bytes = convert_image_to_text(file_bytes)
            if not converted_bytes:
                raise ValueError("No text could be extracted from the image")
            mime_type = "text/plain"

        # PDF to DOCX
        elif input_format == 'pdf' and target_format == 'docx':
            converted_bytes = convert_pdf_to_docx(file_bytes)
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # DOCX to PDF
        elif input_format == 'docx' and target_format == 'pdf':
            converted_bytes = convert_docx_to_pdf(file_bytes)
            mime_type = "application/pdf"

        # Excel to CSV
        elif input_format in ['xlsx', 'xls'] and target_format == 'csv':
            converted_bytes = convert_excel_to_csv(file_bytes)
            mime_type = "text/csv"

        # CSV to Excel
        elif input_format == 'csv' and target_format == 'xlsx':
            converted_bytes = convert_csv_to_excel(file_bytes)
            mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

        else:
            raise ValueError("Unsupported conversion combination")

        if not converted_bytes:
            raise ValueError("Conversion failed: No output generated")

        return converted_bytes, mime_type

    except Exception as e:
        raise Exception(f"Conversion error: {str(e)}")